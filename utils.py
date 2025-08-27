import os
import re
import io
import json
import copy
import requests
from fpdf import FPDF
from flask import request, render_template
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

TEXTBOOKS_API = "https://staticapis.pragament.com/textbooks/allbooks.json"

# -------------------- File I/O --------------------
def read_json(filepath):
    if os.path.exists(filepath):
        with open(filepath, "r") as f:
            return json.load(f)
    return {}

def write_json(data, filepath):
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    with open(filepath, "w") as f:
        json.dump(data, f, indent=4)

# -------------------- Utilities --------------------
def extract_prefix(text):
    match = re.match(r"^([\d\.]+)", text.strip())
    return match.group(1) if match else None

def normalize_chapter_structure(chapter):
    normalized_topics = []
    for topic in chapter.get("topics", []):
        topic_name = topic.get("topic") or topic.get("text")
        subtopics = topic.get("subtopics", [])
        normalized_subtopics = [{"text": sub.get("text")} for sub in subtopics if "text" in sub]
        normalized_topics.append({
            "topic": topic_name,
            "subtopics": normalized_subtopics
        })
    return {
        "number": chapter.get("number"),
        "chapter": chapter.get("chapter"),
        "topics": normalized_topics,
        "reason": chapter.get("reason"),
        "for": chapter.get("for")
    }

# -------------------- API Fetching --------------------
def fetch_textbooks():
    try:
        response = requests.get(TEXTBOOKS_API)
        response.raise_for_status()
        return response.json().get('data', {}).get('getBooks', [])
    except requests.RequestException as e:
        print(f"Error fetching textbooks: {e}")
        return []

def fetch_page_attributes(book_id):
    try:
        url = f"https://staticapis.pragament.com/textbooks/page_attributes/{book_id}.json"
        response = requests.get(url)
        response.raise_for_status()
        return response.json()
    except requests.RequestException as e:
        print(f"Error fetching page attributes for book ID {book_id}: {e}")
        return []

# -------------------- Data Builders --------------------
def extract_topic_and_subtopic_maps(topics, subtopics):
    topic_prefix_map = {
        extract_prefix(t.get('text', '')): {
            "text": t.get("text", ""),
            "subtopics": []
        } for t in topics if extract_prefix(t.get('text', ''))
    }
    subtopic_prefix_map = {
        extract_prefix(s.get('text', '')): s.get('text', '')
        for s in subtopics if extract_prefix(s.get('text', ''))
    }
    for sub_prefix, sub_text in subtopic_prefix_map.items():
        parent_prefix = ".".join(sub_prefix.split('.')[:-1])
        if parent_prefix in topic_prefix_map:
            topic_prefix_map[parent_prefix]['subtopics'].append({"text": sub_text})
    return topic_prefix_map

def build_chapter_structure(data):
    chapters = sorted([item for item in data if item.get("type") == "chapter"], key=lambda x: x.get('order', 0))
    topics = sorted([item for item in data if item.get("type") == "topic"], key=lambda x: x.get('order', 0))
    subtopics = sorted([item for item in data if item.get("type") == "subtopic"], key=lambda x: x.get('order', 0))

    topic_prefix_map = extract_topic_and_subtopic_maps(topics, subtopics)

    final_chapters = []
    chapter_number_name_map = {}
    for idx, ch in enumerate(chapters, start=1):
        chapter_prefix = str(idx)
        chapter_topics = []
        for topic_prefix, topic_data in topic_prefix_map.items():
            if topic_prefix.startswith(chapter_prefix + "."):
                chapter_topics.append({
                    "text": topic_data["text"],
                    "subtopics": topic_data["subtopics"]
                })
        chapter_name = ch.get("text", "")
        final_chapters.append({
            "chapter": chapter_name,
            "number": idx,
            "topics": [
                {
                    "topic": t["text"],
                    "subtopics": t["subtopics"]
                } for t in chapter_topics
            ]
        })
        chapter_number_name_map[idx] = chapter_name
    return final_chapters, chapter_number_name_map

def build_subject_chapter_map(board, class_name, subjects, textbooks):
    subject_chapter_map = {}
    chapter_number_to_name_map = {}
    for subject in subjects:
        matching_book = next(
            (book for book in textbooks
             if book.get('board') == board and str(book.get('class')) == class_name and book.get('subject') == subject),
            None
        )
        if not matching_book:
            print(f"No matching textbook found for subject: {subject}")
            continue
        book_id = matching_book.get("id")
        data = fetch_page_attributes(book_id)
        chapters, chapter_map = build_chapter_structure(data)
        subject_chapter_map[subject] = chapters
        chapter_number_to_name_map[subject] = chapter_map
    return subject_chapter_map, chapter_number_to_name_map

def build_selected_structure(class_name, subjects, chapters, all_chapters_by_subject):
    selected_structure = {
        f"class_{class_name}": {subject: [] for subject in subjects}
    }
    for subject in subjects:
        subject_chapters = all_chapters_by_subject.get(subject, [])
        for ch_name in chapters:
            matched = next((ch for ch in subject_chapters if ch.get("chapter") == ch_name), None)
            if matched and matched not in selected_structure[f"class_{class_name}"][subject]:
                selected_structure[f"class_{class_name}"][subject].append(normalize_chapter_structure(matched))
    return selected_structure

# -------------------- PDF Generators --------------------
def generate_pdf(data, output_pdf):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Generated Question Paper", ln=1, align='C')
    pdf.cell(200, 10, txt=f"Board: {data.get('board', 'N/A')}, Class: {data.get('class', 'N/A')}, Subject: {', '.join(data.get('subject', []))}", ln=2, align='C')
    pdf.ln(10)
    for idx, q in enumerate(data.get("questions", []), start=1):
        pdf.multi_cell(0, 10, txt=f"{idx}. {q.get('question')}")
        pdf.ln(3)
        for i, opt in enumerate(q.get("options", []), start=1):
            pdf.cell(0, 10, txt=f"({chr(64+i)}) {opt}", ln=1)
        pdf.cell(0, 10, txt=f"Correct Answer: {q.get('correct_answer')}", ln=1)
        pdf.ln(5)
    pdf.output(output_pdf)

def generate_docx(data, output_docx):
    """Generate a Word (.docx) file for the question paper."""
    doc = Document()
    doc.add_heading('Generated Question Paper', 0)
    board = data.get('board', 'N/A')
    class_name = data.get('class', 'N/A')
    subject = ', '.join(data.get('subject', [])) if isinstance(data.get('subject'), list) else data.get('subject', 'N/A')
    doc.add_paragraph(f"Board: {board}, Class: {class_name}, Subject: {subject}")
    doc.add_paragraph("")
    questions = data.get('questions', [])
    for idx, q in enumerate(questions, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {q.get('question')}")
        run.font.size = Pt(12)
        options = q.get('options', [])
        for i, opt in enumerate(options, start=1):
            p = doc.add_paragraph(f"({chr(64+i)}) {opt}", style='List Bullet')
            p.paragraph_format.left_indent = Pt(24)
        correct = q.get('correct_option')
        if correct:
            correct_text = options[correct-1] if 1 <= correct <= len(options) else "N/A"
            doc.add_paragraph(f"Correct Answer: {correct_text}")
        doc.add_paragraph("")
    doc.save(output_docx)
def generate_prerequisite_pdf(tree):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin_left, margin_right, margin_top, margin_bottom = 30, 30, 40, 40
    y = height - margin_top
    line_height = 22
    max_indent = 80
    sidebar_width = 5
    background_colors = [colors.whitesmoke, colors.lightgrey]
    sidebar_colors = [colors.red, colors.orange, colors.green, colors.cadetblue, colors.purple, colors.brown, colors.teal]
    line_counter = 0

    def draw_text_block(text, level, font="Helvetica", font_size=11):
        nonlocal y, line_counter
        if y < margin_bottom + line_height:
            c.showPage()
            y = height - margin_top
            line_counter = 0
        indent = min(level * 20, max_indent)
        x_pos = margin_left + indent + sidebar_width + 5
        bg_color = background_colors[line_counter % 2]
        c.setFillColor(bg_color)
        c.rect(margin_left, y - line_height + 4, width - margin_left - margin_right, line_height, fill=1, stroke=0)
        sidebar_color = sidebar_colors[level % len(sidebar_colors)]
        c.setFillColor(sidebar_color)
        c.rect(margin_left, y - line_height + 4, sidebar_width, line_height, fill=1, stroke=0)
        c.setFillColor(colors.black)
        c.setFont(font, font_size)
        c.drawString(x_pos, y, text)
        y -= line_height
        line_counter += 1

    def draw_chapters(chapters, level=0):
        for chapter in chapters:
            chapter_text = f"{chapter['chapter']} (Chapter {chapter['number']}, {chapter['class']})"
            draw_text_block(chapter_text, level, font="Helvetica-Bold", font_size=12)
            if "reason" in chapter:
                draw_text_block(f"Reason: {chapter['reason']}", level + 1, font="Helvetica-Oblique", font_size=10)
            if chapter.get("prerequisites"):
                draw_chapters(chapter["prerequisites"], level + 1)

    for class_key, subjects in tree.items():
        for subject, chapters in subjects.items():
            draw_text_block(f"{subject} - {class_key}", 0, font="Helvetica-Bold", font_size=14)
            draw_chapters(chapters, level=1)
            y -= line_height // 2

    c.save()
    buffer.seek(0)
    return buffer

# === File I/O ===
def save_json(data, path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w") as f:
        json.dump(data, f, indent=4)

def load_json(path):
    if os.path.exists(path):
        with open(path, "r") as f:
            return json.load(f)
    return {}

# === Main Handler (Top Level Usage) ===
def handle_final_level(level, class_name, subjects, selected_structure_path):
    render_path = f"structured_data/prereq_render_items_level_{level - 1}.json"
    render_items = load_json(render_path)
    selected_structure = load_json(selected_structure_path)

    selected_combined = request.form.getlist("selected_prereq_combined")
    selected_topics = request.form.getlist("selected_prereq_topic")
    selected_subtopics = request.form.getlist("selected_prereq_subtopic")

    selected_ids, selected_chapters = [], []
    for item in selected_combined:
        try:
            id_, chapter = item.split("|||", 1)
            selected_ids.append(id_)
            selected_chapters.append(chapter)
        except ValueError:
            print(f"⚠️ Skipped invalid combined item: {item}")

    id_map = {item["id"]: item for item in render_items}
    selected_items = [id_map[i] for i in selected_ids if i in id_map]

    class_key = f"class_{int(class_name) - level + 1}"
    selected_structure.setdefault(class_key, {})
    for subject in subjects:
        selected_structure[class_key].setdefault(subject, [])

    previous_level_data = load_json(f"structured_data/previous_year_depth_{level - 1}.json")

    for item in selected_items:
        subject = item["subject"]
        chapter_name = item["chapter"]
        matched = next((ch for ch in previous_level_data.get(subject, []) if ch["chapter"] == chapter_name), None)
        if not matched:
            print(f"❌ No match found in previous level for: {chapter_name}")
            continue
        chapter_obj = normalize_chapter_structure(matched)
        chapter_obj["for"] = item.get("for")
        chapter_obj["reason"] = item.get("reason")
        if not any(existing["chapter"] == chapter_obj["chapter"] and existing.get("for") == chapter_obj.get("for")
                   for existing in selected_structure[class_key][subject]):
            print(f"➕ Adding Chapter: {chapter_obj}")
            selected_structure[class_key][subject].append(chapter_obj)

    # ✅ Match by topic/subtopic
    for subject in previous_level_data:
        for chapter in previous_level_data[subject]:
            for topic in chapter.get("topics", []):
                if topic.get("topic") in selected_topics or any(
                        sub.get("text") in selected_subtopics for sub in topic.get("subtopics", [])):
                    chapter_obj = normalize_chapter_structure(chapter)
                    selected_structure[class_key][subject].append(chapter_obj)
                    break

    # ✅ Inject reasons
    reason_map = {
        (item.get("subject", "").strip().lower(), item.get("chapter", "").strip().lower(),
         item.get("for", "").strip().lower()): {
            "reason": item.get("reason"),
            "for": item.get("for")
        } for item in render_items
    }

    for c_key, subjects_data in selected_structure.items():
        for subject, chapter_list in subjects_data.items():
            subj_key = subject.strip().lower()
            for chapter in chapter_list:
                chap_key = (chapter.get("chapter") or "").strip().lower()
                for_key = (chapter.get("for") or "").strip().lower()
                reason = reason_map.get((subj_key, chap_key, for_key))
                if reason:
                    chapter["reason"] = reason["reason"]
                    chapter["for"] = reason["for"]

    save_json(selected_structure, selected_structure_path)

    tree = build_prerequisite_tree(selected_structure)
    save_json(tree, "structured_data/prerequisite_tree.json")
    return render_template("next_step.html", tree_json=tree)

# === Tree Builders ===
def build_prerequisite_tree(selected_structure):
    sorted_classes = sorted(selected_structure.keys(), key=lambda k: int(k.split('_')[1]), reverse=True)
    if not sorted_classes: return {}

    top_class_key = sorted_classes[0]
    result = copy.deepcopy(selected_structure[top_class_key])
    for subject, chapters in result.items():
        for chapter in chapters:
            chapter["class"] = top_class_key

    def attach(subject, chapter_name, cur_idx, visited=None):
        visited = visited or set()
        if chapter_name in visited: return []
        visited.add(chapter_name)

        prereqs = []
        for lower_idx in range(cur_idx + 1, len(sorted_classes)):
            class_key = sorted_classes[lower_idx]
            chapters = selected_structure.get(class_key, {}).get(subject, [])
            for ch in chapters:
                if ch.get("chapter") in visited: continue
                if ch.get("for") == chapter_name or not ch.get("for"):
                    ch_copy = copy.deepcopy(ch)
                    ch_copy.pop("for", None)
                    ch_copy["class"] = class_key
                    ch_copy["prerequisites"] = attach(subject, ch.get("chapter"), lower_idx, visited.copy())
                    prereqs.append(ch_copy)
        return prereqs

    for subject, chapters in result.items():
        for chapter in chapters:
            chapter["prerequisites"] = attach(subject, chapter.get("chapter"), 0)

    return {top_class_key: result}

def build_prerequisite_tree_minimal(selected_structure):
    sorted_classes = sorted(selected_structure.keys(), key=lambda k: int(k.split('_')[1]), reverse=True)
    if not sorted_classes: return {}

    top_class_key = sorted_classes[0]
    result = {}

    def minimal(ch, class_key):
        obj = {
            "number": ch.get("number"),
            "chapter": ch.get("chapter"),
            "class": class_key
        }
        if ch.get("reason"): obj["reason"] = ch["reason"]
        return obj

    def attach(subject, chapter_name, cur_idx, visited=None):
        visited = visited or set()
        if (subject, chapter_name) in visited: return []
        visited.add((subject, chapter_name))

        prereqs = []
        for lower_idx in range(cur_idx + 1, len(sorted_classes)):
            class_key = sorted_classes[lower_idx]
            chapters = selected_structure.get(class_key, {}).get(subject, [])
            for ch in chapters:
                if ch.get("for", "").strip() == chapter_name:
                    ch_entry = minimal(ch, class_key)
                    ch_entry["prerequisites"] = attach(subject, ch.get("chapter"), lower_idx, visited)
                    prereqs.append(ch_entry)
        return prereqs

    result[top_class_key] = {}
    for subject, chapters in selected_structure[top_class_key].items():
        result[top_class_key][subject] = []
        for ch in chapters:
            ch_entry = minimal(ch, top_class_key)
            ch_entry["prerequisites"] = attach(subject, ch.get("chapter"), 0)
            result[top_class_key][subject].append(ch_entry)

    return result

# === Fetchers ===
def fetch_structured_previous_year_content(board, class_name, subjects, depth=1, max_depth=3):
    if depth > max_depth: return {}

    previous_class = str(int(class_name) - 1)
    try:
        response = requests.get(TEXTBOOKS_API)
        response.raise_for_status()
        textbooks = response.json().get('data', {}).get('getBooks', [])
    except requests.RequestException:
        return {}

    full_structure = {}

    for subject in subjects:
        book = next((b for b in textbooks if str(b.get('class')) == previous_class and b.get('board') == board and b.get('subject') == subject), None)
        if not book: continue

        try:
            response = requests.get(f"https://staticapis.pragament.com/textbooks/page_attributes/{book['id']}.json")
            response.raise_for_status()
            data = response.json()
        except requests.RequestException:
            continue

        chapters = sorted([i for i in data if i.get("type") == "chapter"], key=lambda x: x.get("order", 0))
        topics = sorted([i for i in data if i.get("type") == "topic"], key=lambda x: x.get("order", 0))
        subtopics = sorted([i for i in data if i.get("type") == "subtopic"], key=lambda x: x.get("order", 0))

        topic_prefix_map = {}
        for t in topics:
            prefix = extract_prefix(t.get("text", ""))
            if prefix: topic_prefix_map[prefix] = {"text": t["text"], "subtopics": []}

        for s in subtopics:
            sub_prefix = extract_prefix(s.get("text", ""))
            parent = ".".join(sub_prefix.split('.')[:-1]) if sub_prefix else None
            if parent in topic_prefix_map:
                topic_prefix_map[parent]["subtopics"].append({"text": s["text"]})

        chapter_data = []
        for idx, ch in enumerate(chapters, start=1):
            chapter_topics = []
            for tp, data in topic_prefix_map.items():
                if tp.startswith(f"{idx}."):
                    chapter_topics.append({"text": data["text"], "subtopics": data["subtopics"]})
            chapter_data.append({"chapter": ch["text"], "number": idx, "topics": chapter_topics})

        full_structure[subject] = chapter_data

    save_json(full_structure, f"structured_data/previous_year_depth_{depth}.json")
    return full_structure

# === Prompt Builder ===
def build_prompt(subject, chapter_name, previous_year_chapters):
    return f"""
    You are an academic AI assistant helping to identify prerequisite chapters.

    Context:
    - The user has selected a specific chapter from a current year's syllabus.
    - You are also given the chapter list from the previous year's syllabus for the same subject and board.

    Instructions:
    1. Identify only those prerequisite chapters that are clearly and directly related.
    2. Avoid abstract or general background prerequisites.
    3. All suggested prerequisites must come from the previous year's chapter list.

    Output Format:
    {{
    "prerequisites": {{
        "{subject}": [
        {{
            "number": 1,
            "chapter": "Exact Chapter Name",
            "reason": "Why this chapter is needed",
            "for": "{chapter_name}"
        }}
        ]
    }}
    }}

    Selected Chapter:
    {json.dumps([{"chapter": chapter_name}], indent=2)}

    Previous Year Chapters:
    {json.dumps(previous_year_chapters, indent=2)}
    """